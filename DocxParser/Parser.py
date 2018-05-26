from DocxParser.source.TitleClass import Title
import docx
import zipfile

class ParserProgram:
    def __init__(self):
        self.titlelist = []  # titlelist, which contain title classes
        self.tempParent1 = None  # index==1인  클래스의 부모를 가리키는 포인터
        self.tempParent2 = None  # index==2인 클래스의 부모를 가리키는 포인터
        self.tempParent3 = None  # index=3인 클래스의 부모를 가리키는 포인터
        self.isBold = True
        self.isItalic = True
        self.isHighlight = True
        self.isUnderLine = True
        self.doc=None
    def set_parserobject(self,Bold,Italic,Highlight,UnderLine):
        self.isBold = Bold
        self.isItalic = Italic
        self.isHighlight = Highlight
        self.isUnderLine = UnderLine
    def indentcount(self,aparagraph):
        i = 0
        icount = aparagraph.paragraph_format.left_indent
        if (icount is not None):
            i = icount / 127000
        return i
    def load_file(self,filename):
        if len(self.titlelist) is not 0:
            while len(self.titlelist) is not 0:
                del(self.titlelist[0])

        self.doc = docx.Document(filename)
        paranum = len(self.doc.paragraphs)  # Number of paragraphs in docx file
        for i in range(paranum):
            icount = self.indentcount(self.doc.paragraphs[i])
            if self.doc.paragraphs[i].text != '':
                if icount == 0:
                    atitle = Title()
                    atitle.paraindex=i
                    atitle.index = 0
                    atitle.title = self.doc.paragraphs[i].text
                    atitle.content.append(self.doc.paragraphs[i].text)
                    atitle.TreeConstructor(None)
                    self.tempParent1 = atitle
                    self.titlelist.append(atitle)
                elif icount == 1:
                    atitle = Title()
                    atitle.paraindex=i
                    atitle.index = 1
                    atitle.title = self.doc.paragraphs[i].text
                    atitle.content.append(self.doc.paragraphs[i].text)
                    atitle.TreeConstructor(self.tempParent1)
                    self.tempParent2 = atitle
                    self.titlelist.append(atitle)
                elif icount == 2:
                    atitle = Title()
                    atitle.paraindex=i
                    atitle.index = 2
                    atitle.title = self.doc.paragraphs[i].text
                    atitle.content.append(self.doc.paragraphs[i].text)
                    atitle.TreeConstructor(self.tempParent2)
                    self.tempParent3 = atitle
                    self.titlelist.append(atitle)
                else:
                    atitle=Title()
                    atitle.paraindex=i
                    atitle.index=3
                    atitle.TreeConstructor(self.tempParent3)
                    self.titlelist.append(atitle)

        listlen = len(self.titlelist)
        dellist = []
        for a in range(listlen):
            if len(self.titlelist[a].child) == 0:
                parent = self.titlelist[a].parent
                runlen = self.doc.paragraphs[self.titlelist[a].paraindex].runs.__len__()
                transition = -1
                BoldRunSet = ""
                HighlightRunSet = ""
                UnderLineRunSet = ""
                ItalicRunSet = ""
                for j in range(runlen):
                    if self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].bold == True and self.isBold:
                        if transition == 0 or transition == -1:
                            BoldRunSet = BoldRunSet + self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].text
                        elif transition == 1:
                            if ItalicRunSet != "":
                                self.titlelist[a].parent.content.append(ItalicRunSet)
                            ItalicRunSet = ""
                        elif transition == 2:
                            if HighlightRunSet != "":
                                self.titlelist[a].parent.content.append(HighlightRunSet)
                            HighlightRunSet = ""
                        elif transition == 3:
                            if UnderLineRunSet != "":
                                self.titlelist[a].parent.content.append(UnderLineRunSet)
                            UnderLineRunSet = ""
                        transition = 0
                    elif self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].italic == True and self.isItalic:
                        if transition == 0:
                            if BoldRunSet != "":
                                self.titlelist[a].parent.content.append(BoldRunSet)
                            BoldRunSet = ""
                        elif transition == 1 or transition == -1:
                            ItalicRunSet = ItalicRunSet + self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].text
                        elif transition == 2:
                            if HighlightRunSet != "":
                                self.titlelist[a].parent.content.append(HighlightRunSet)
                            HighlightRunSet = ""
                        elif transition == 3:
                            if UnderLineRunSet != "":
                                self.titlelist[a].parent.content.append(UnderLineRunSet)
                            UnderLineRunSet = ""
                        transition = 1
                    elif self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].font.highlight_color is not None and self.isHighlight:
                        if transition == 0:
                            if BoldRunSet != "":
                                self.titlelist[a].parent.content.append(BoldRunSet)
                            BoldRunSet = ""
                        elif transition == 1:
                            if ItalicRunSet != "":
                                self.titlelist[a].parent.content.append(ItalicRunSet)
                            ItalicRunSet = ""
                        elif transition == 2 or transition == -1:
                            HighlightRunSet = HighlightRunSet + self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].text
                        elif transition == 3:
                            if UnderLineRunSet != "":
                                self.titlelist[a].parent.content.append(UnderLineRunSet)
                            UnderLineRunSet = ""
                        transition = 2
                    elif self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].underline == True and self.isUnderLine:
                        if transition == 0:
                            if BoldRunSet != "":
                                self.titlelist[a].parent.content.append(BoldRunSet)
                            BoldRunSet = ""
                        elif transition == 1:
                            if ItalicRunSet != "":
                                self.titlelist[a].parent.content.append(ItalicRunSet)
                            ItalicRunSet = ""
                        elif transition == 2:
                            if HighlightRunSet != "":
                                self.titlelist[a].parent.content.append(HighlightRunSet)
                            HighlightRunSet = ""
                        elif transition == 3 or transition == -1:
                            UnderLineRunSet = UnderLineRunSet + self.doc.paragraphs[self.titlelist[a].paraindex].runs[j].text
                        transition = 3
                    else:
                        if BoldRunSet != "":
                            self.titlelist[a].parent.content.append(BoldRunSet)
                            BoldRunSet = ""
                        elif ItalicRunSet != "":
                            self.titlelist[a].parent.content.append(ItalicRunSet)
                            ItalicRunSet = ""
                        elif HighlightRunSet != "":
                            self.titlelist[a].parent.content.append(HighlightRunSet)
                            HighlightRunSet = ""
                        elif UnderLineRunSet != "":
                            self.titlelist[a].parent.content.append(UnderLineRunSet)
                            UnderLineRunSet = ""
                        transition = -1
                if j == runlen - 1:
                    if BoldRunSet != "":
                        self.titlelist[a].parent.content.append(BoldRunSet)
                        BoldRunSet = ""
                    elif ItalicRunSet != "":
                        self.titlelist[a].parent.content.append(ItalicRunSet)
                        ItalicRunSet = ""
                    elif HighlightRunSet != "":
                        self.titlelist[a].parent.content.append(HighlightRunSet)
                        HighlightRunSet = ""
                    elif UnderLineRunSet != "":
                        self.titlelist[a].parent.content.append(UnderLineRunSet)
                        UnderLineRunSet = ""
                    transition = -1

                contentlen = len(self.titlelist[a].content)
                for c in range(contentlen):
                    if c!=0:
                        parent.content.append(self.titlelist[a].content[c])
                dellist.append(a)
        dellistlen = len(dellist)
        for b in range(dellistlen):
            del (self.titlelist[dellist[b] - b])
        print("First phase!")
        listlen=len(self.titlelist)
        for a in range(listlen):
            print(self.titlelist[a].title)
            print(self.titlelist[a].content)


    def load_image(self, filename):
        z = zipfile.ZipFile(filename)
        all_files = z.namelist()
        images = filter(lambda x: x.startswith('word/media/'), all_files)
        for image in images:
            print(image)
            image1 = z.open(image).read()
            f = open(image[11:], 'wb')
            f.write(image1)
            f.close()
        z.close()

